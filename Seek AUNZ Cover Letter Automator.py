#importing modules
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException
import time

#opening firefox and going to Google homepage
browser = webdriver.Firefox()
browser.get('https://www.google.co.nz')


#going to seek.co.nz's website
search_elem = browser.find_element(By.CSS_SELECTOR, '#APjFqb')
search_elem.click()
time.sleep(3)
search_elem.send_keys('seek.co.nz')
search_elem.submit()
time.sleep(5)
seek_elem = browser.find_element(By.CSS_SELECTOR, '.eKjLze > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > span:nth-child(1) > a:nth-child(1) > h3:nth-child(2)')
seek_elem.click()
#congratulations, you have gained access to seeks homepage website


#now signing into you're account
sign_in_elem = browser.find_element(By.CSS_SELECTOR, ' div.a1msqiga:nth-child(3) > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > div:nth-child(2) > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > a:nth-child(1)')
sign_in_elem.click()
time.sleep(5)
email_elem = browser.find_element(By.CSS_SELECTOR, '#emailAddress')
email_elem.click()
email_elem.send_keys('insert you're username')
pass_elem = browser.find_element(By.CSS_SELECTOR, '#password')
pass_elem.click()  
pass_elem.send_keys('insert you're password')
pass_elem.submit()
time.sleep(10)

#now finding you're saved jobs list
savedjobs = browser.find_element(By.CSS_SELECTOR, '.a1msqi5c > div:nth-child(2) > div:nth-child(1) > nav:nth-child(1) > div:nth-child(2) > div:nth-child(1) > div:nth-child(1) > div:nth-child(1) > a:nth-child(1)')
savedjobs.click()
time.sleep(5)

job_details = []

expired_found = False

while not expired_found:
#Obtaing links to all saved jobs in list
# Get all job links
    job_elements = browser.find_elements(By.CSS_SELECTOR, 'a.g08t882.g08t884.g08t887.aoyz142.d3eu8q0.d3eu8qf._1igc8rlh')
    links = []
    for job_element in job_elements:
        link = job_element.get_attribute("href")
        if not link.endswith('/apply') and not link.endswith('apply/linkout') and not'/expiredjob' in link:
            links.append(link)

# Open each job link in a new browser tab
    main_window = browser.current_window_handle

    for link in links:
# Open a new tab
        browser.execute_script(f"window.open('{link}');")
        time.sleep(3)
        
# Switch to the new tab    
        new_window = [window for window in browser.window_handles if window != main_window][0]
        browser.switch_to.window(new_window)
      
        
# Scrape the job title, recruiter, and category
        title = browser.find_element(By.CSS_SELECTOR, 'h1._1wkzzau0.a1msqi4y.lnocuo0.lnocuol._1d0g9qk4.lnocuop.lnocuo21').text
        time.sleep(3)
        avoid_characters = ["\\", "/", ":", "*", "?", "<", ">", "|", "–"]
        for char in avoid_characters:
            if char in title:
                title = title.replace(char, "-")
        recruiter = browser.find_element(By.CSS_SELECTOR, 'span._1wkzzau0.a1msqi4y.lnocuo0.lnocuo1.lnocuo21._1d0g9qk4.lnocuod').text
        for char in avoid_characters: 
            if char in recruiter:
                recruiter = recruiter.replace(char, "-")
            time.sleep(5)

# Finding the job classification field needed
        finding_category = browser.find_elements(By.CSS_SELECTOR, 'span._1wkzzau0.a1msqi4y.a1msqir')
        category = ""
        for unwanted in finding_category:
            text = unwanted.text
            if "Expired" and "(AUD)" in text:
                continue
            elif "(" in text and ")" in text:
                category = text

        time.sleep(2)

# Storing job information
        job_details.append({
        'title': title,
        'recruiter': recruiter,
        'category': category,
        })

# Close the tab and switch back to the main window
        browser.close()
        browser.switch_to.window(main_window)
        time.sleep(3)

#trying to find the next page element if not found.
    try:
        next_page_button = browser.find_element(By.CSS_SELECTOR, "[data-automation='pager-next']")
        next_page_button.click()
        time.sleep(3)
    except NoSuchElementException:
        break
    except Exception as e:
        print(f"error found: {e}")
#congratulations all job details required have been extracted and saved. 

print('completed obtaining all job details\n')
print(job_details)

#specifying which cover letter template to use for which job catgeory and what sections of each template to replace for a particular job. 
import docx
for jobs in job_details:
#assigning templates to variables. 
    finance_cl_template = docx.Document('D:\\Desktop\\cl automator\\finance cl template.docx')
    retail_and_marketing_cl_template = docx.Document('D:\\Desktop\\cl automator\\retail and marketing cl template.docx')
    data_and_tech_template = docx.Document('d:\\desktop\\cl automator\\data and tech template.docx')
    manufacturing_transport_and_logistics_cl_template = docx.Document('d:\\desktop\\cl automator\\manufacturing transport and logistics cl template.docx')
    trades_and_services_cl_template = docx.Document('d:\\desktop\\cl automator\\trades and services cl template.docx')
#with the below specifying sections to amend for each cover letter template. 
    if "(Banking & Financial Services)" in jobs['category'] or "(Accounting)" in jobs['category'] or "(Consulting & Strategy)" in jobs['category'] or "(Sales)" in jobs['category']:
        finance_cl_template.paragraphs[2].text = finance_cl_template.paragraphs[2].text.replace('Payments Officer – New Market', jobs['title'])
        finance_cl_template.paragraphs[3].text = finance_cl_template.paragraphs[3].text.replace('Payments Officer', jobs['title']).replace('ACC', jobs['recruiter'])
        finance_cl_template.paragraphs[12].text = finance_cl_template.paragraphs[12].text.replace('ACC', jobs['recruiter'])
        finance_cl_template.save(f'D:\\Desktop\\cl automator\\{jobs["recruiter"]} - {jobs["title"]} Cover Letter.docx')
    elif "(Information & Communication Technology)" in jobs['category']:
        data_and_tech_template.paragraphs[1].text = data_and_tech_template.paragraphs[1].text.replace('Technical Specialist', jobs['title']) 
        data_and_tech_template.paragraphs[9].text = data_and_tech_template.paragraphs[9].text.replace('Kinetic Recruitment', jobs['recruiter'])
        data_and_tech_template.save(f'D:\\Desktop\\cl automator\\{jobs["recruiter"]} - {jobs["title"]} Cover Letter.docx')
    elif "(Administration & Office Support)" in jobs['category']:
        retail_and_marketing_cl_template.paragraphs[2].text = retail_and_marketing_cl_template.paragraphs[2].text.replace('Spark Pukekohe', jobs['title'])
        retail_and_marketing_cl_template.paragraphs[4].text = retail_and_marketing_cl_template.paragraphs[4].text.replace('Spark Technician', jobs['title'])
        retail_and_marketing_cl_template.paragraphs[13].text = retail_and_marketing_cl_template.paragraphs[13].text.replace('Spark New Zealand Trading Limited', jobs['recruiter'])
        retail_and_marketing_cl_template.save(f'D:\\Desktop\\cl automator\\{jobs["recruiter"]} - {jobs["title"]} Cover Letter.docx')
    elif "(Manufacturing, Transport & Logistics)" in jobs['category']:
        manufacturing_transport_and_logistics_cl_template.paragraphs[2].text = manufacturing_transport_and_logistics_cl_template.paragraphs[2].text.replace('Job Title', jobs['title'])
        manufacturing_transport_and_logistics_cl_template.paragraphs[3].text = manufacturing_transport_and_logistics_cl_template.paragraphs[3].text.replace('Job Title', jobs['title'])
        manufacturing_transport_and_logistics_cl_template.paragraphs[3].text = manufacturing_transport_and_logistics_cl_template.paragraphs[3].text.replace('Company Name', jobs['recruiter'])
        manufacturing_transport_and_logistics_cl_template.paragraphs[5].text = manufacturing_transport_and_logistics_cl_template.paragraphs[5].text.replace('Company Name', jobs['recruiter'])
        manufacturing_transport_and_logistics_cl_template.save(f'D:\\Desktop\\cl automator\\{jobs["recruiter"]} - {jobs["title"]} Cover Letter.docx')
    elif "(Trades & Services)" in jobs['category']:
        trades_and_services_cl_template.paragraphs[2].text = trades_and_services_cl_template.paragraphs[2].text.replace('Job Title', jobs['title'])
        trades_and_services_cl_template.paragraphs[3].text = trades_and_services_cl_template.paragraphs[3].text.replace('Job Title', jobs['title'])
        trades_and_services_cl_template.paragraphs[3].text = trades_and_services_cl_template.paragraphs[3].text.replace('Company Name', jobs['recruiter'])
        trades_and_services_cl_template.paragraphs[5].text = trades_and_services_cl_template.paragraphs[5].text.replace('Company Name', jobs['recruiter'])
        trades_and_services_cl_template.save(f'D:\\Desktop\\cl automator\\{jobs["recruiter"]} - {jobs["title"]} Cover Letter.docx')
    else: #any (category in jobs['category'] for category in jobs['category']):
        retail_and_marketing_cl_template.paragraphs[2].text = retail_and_marketing_cl_template.paragraphs[2].text.replace('Spark Pukekohe', jobs['title'])
        retail_and_marketing_cl_template.paragraphs[4].text = retail_and_marketing_cl_template.paragraphs[4].text.replace('Spark Technician', jobs['title'])
        retail_and_marketing_cl_template.paragraphs[13].text = retail_and_marketing_cl_template.paragraphs[13].text.replace('Spark New Zealand Trading Limited', jobs['recruiter'])
        retail_and_marketing_cl_template.save(f'D:\\Desktop\\cl automator\\{jobs["recruiter"]} - {jobs["title"]} Cover Letter.docx')

print('\nall job cover letters have been created that were not expired')
    #finance_cl_template = docx.Document('D:\\Desktop\\cl automator\\finance cl template.docx')
